import io
import time
try:
    import pandas as pd
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pd = None
    Document = None

def generate_excel_bytes(biddings):
    if not pd:
        return b""
    data = []
    for b in biddings:
        # 处理供应商名称显示逻辑，与前端保持一致
        supplier_name = '未知'
        if b.meta_info and isinstance(b.meta_info, dict):
            if b.meta_info.get('supplier') and b.meta_info.get('supplier') != '未知':
                supplier_name = b.meta_info.get('supplier')
        
        data.append({
            "标题": b.title,
            "网站来源": b.source_website,
            "公告类型": b.notice_type,
            "发布日期": b.publish_date.strftime('%Y-%m-%d') if b.publish_date else "",
            "供应商名称": supplier_name,
            "分类": b.category,
            "简报": b.content_abstract,
            "竞对分析": b.opportunity_analysis,
            "原文链接": b.source_url
        })
    df = pd.DataFrame(data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='竞对清单')
    return output.getvalue()

def generate_word_bytes(biddings):
    if not Document:
        return b""
    doc = Document()
    title = doc.add_heading('竞对雷达 - 智能分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"报告说明: 本报告包含本次抓取到的最新竞对动态（共 {len(biddings)} 条）。")
    for i, b in enumerate(biddings, 1):
        doc.add_heading(f"{i}. {b.title}", level=1)
        p_meta = doc.add_paragraph()
        
        # 处理供应商名称显示逻辑
        supplier_name = '未知'
        if b.meta_info and isinstance(b.meta_info, dict):
            if b.meta_info.get('supplier') and b.meta_info.get('supplier') != '未知':
                supplier_name = b.meta_info.get('supplier')
                
        p_meta.add_run(f"来源: {b.source_website} | 类型: {b.notice_type} | 发布日期: {b.publish_date.strftime('%Y-%m-%d') if b.publish_date else '未知'}\n").bold = True
        p_meta.add_run(f"中标供应商: {supplier_name}").bold = True
        doc.add_heading("项目简报", level=2)
        doc.add_paragraph(b.content_abstract or "暂无")
        doc.add_heading("竞对分析", level=2)
        doc.add_paragraph(b.opportunity_analysis or "暂无")
        doc.add_paragraph(f"原文链接: {b.source_url}")
        doc.add_page_break()
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def generate_email_html(biddings):
    html = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
            .header {{ background-color: #07c160; color: white; padding: 15px; text-align: center; border-radius: 8px; }}
            .bidding-item {{ border: 1px solid #eee; margin: 15px 0; padding: 15px; border-radius: 8px; box-shadow: 0 2px 5px rgba(0,0,0,0.05); }}
            .title {{ color: #2c3e50; font-size: 18px; margin-bottom: 10px; border-bottom: 2px solid #07c160; padding-bottom: 5px; }}
            .meta {{ font-size: 13px; color: #666; margin-bottom: 10px; }}
            .section-title {{ font-weight: bold; color: #07c160; margin-top: 10px; }}
            .content {{ margin-top: 5px; font-size: 14px; background: #f9f9f9; padding: 10px; border-radius: 5px; }}
            .link {{ margin-top: 10px; font-size: 13px; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h2>竞对雷达 - 智能分析报告</h2>
            <p>生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
    """
    
    if not biddings:
        html += """
        <div style="text-align: center; padding: 40px 20px; background: #f9f9f9; border-radius: 8px; margin-top: 20px;">
            <h3 style="color: #666;">本次自动抓取未发现新的竞对动态</h3>
            <p style="color: #999; font-size: 14px;">(可能原因: 网站暂无新公告、或者公告内容未命中竞对)</p>
        </div>
        """
    else:
        html += f"<p>本次共提取到 <b>{len(biddings)}</b> 条竞对动态，详细清单请查看附件 Excel 及 Word 报告。</p>"
        for i, b in enumerate(biddings[:20], 1):  # 邮件正文最多显示20条，避免过长
            html += f"""
            <div class="bidding-item">
                <div class="title">{i}. {b.title}</div>
                <div class="meta">
                    来源: {b.source_website} | 发布日期: {b.publish_date.strftime('%Y-%m-%d') if b.publish_date else '未知'}
                </div>
                <div class="section-title">项目简报</div>
                <div class="content">{b.content_abstract or "暂无"}</div>
                <div class="section-title">竞对分析</div>
                <div class="content">{b.opportunity_analysis or "暂无"}</div>
                <div class="link"><a href="{b.source_url}" target="_blank">查看原文公告</a></div>
            </div>
            """
        if len(biddings) > 20:
            html += f"<p>...还有 {len(biddings) - 20} 条动态，请在附件中查看完整报告。</p>"
    
    html += """
        <p style="text-align: center; color: #888; font-size: 12px; margin-top: 30px;">
            本报告由竞对雷达自动生成
        </p>
    </body>
    </html>
    """
    return html
